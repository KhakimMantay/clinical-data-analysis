import os
import csv
import re
import hmac
import hashlib
from pathlib import Path
from datetime import datetime
from collections import Counter

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
IN_CSV = PROJECT_ROOT / "file_index.csv"
OUT_CSV = PROJECT_ROOT / "file_index_with_dates.csv"

PROGRESS_EVERY = 200

MIN_YEAR = 2019
MAX_YEAR = 2026

TWO_DIGIT_YEAR_CUTOFF = 80  # 00..79 => 20xx, 80..99 => 19xx
PATIENT_SECRET_ENV = "PATIENT_ID_SECRET"

FIO_RE = re.compile(
    r"(?:Ф\.?\s*И\.?\s*О\.?|ФИО\s*пациента|ФИО|Фамилия\s*Имя\s*Отчество|Пациент(?:ка)?)\s*[:\-]\s*(.+)",
    re.IGNORECASE,
)

DOB_RE = re.compile(
    r"(?:Дата\s*рождения|Число\s*,\s*месяц\s*,\s*год\s*рождения|Туған\s+күні\s*\(Дата\s+рождения\))\s*[:\-]?\s*"
    r"([0-3]?\d[.\-/][01]?\d[.\-/]\d{4})\s*(?:г\.?\s*р\.?|гр\.?)?",
    re.IGNORECASE,
)

DOB_GR_RE = re.compile(
    r"\b([0-3]?\d[.\-/][01]?\d[.\-/]\d{4})\s*(?:г\.?\s*р\.?|гр\.?)\b",
    re.IGNORECASE,
)

PATIENT_NAME_BLOCK_RE = re.compile(
    r"(?:ФИО\s*пациента|Т\.\s*А\.\s*Ө\.\s*\(Ф\.\s*И\.\s*О\.\)|Пациент(?:ка)?)\s*[:\-]\s*\[?\s*([^\]\n]+?)\s*\]?(?:\n|$)",
    re.IGNORECASE,
)

PATIENT_DOB_BLOCK_RE = re.compile(
    r"(?:Дата\s*рождения|Число\s*,\s*месяц\s*,\s*год\s*рождения|Туған\s+күні\s*\(Дата\s+рождения\))\s*[:\-]?\s*\[?\s*"
    r"([0-3]?\d[.\-/][01]?\d[.\-/]\d{4})",
    re.IGNORECASE,
)

MULTISPACE_RE = re.compile(r"\s+")
NON_LETTER_RE = re.compile(r"[^a-zа-яё-]+", re.IGNORECASE)

FILENAME_STOPWORDS = (
    "перв", "первич", "повт", "повтор", "узи", "онк", "шабл", "шаблон",
    "прием", "приём", "протокол", "копия", "new", "новый",
)

DATE_GENERIC_RE = re.compile(
    r"\b([0-3]?\d)\s*[./,\-]\s*([01]?\d)\s*[./,\-]\s*((?:19|20)\d{2}|\d{2})(?=\D|$)",
    re.IGNORECASE,
)
RU_MONTHS = {
    "января": 1, "февраля": 2, "марта": 3, "апреля": 4, "мая": 5, "июня": 6,
    "июля": 7, "августа": 8, "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12,
}
RU_MONTH_RE = re.compile(
    r"\b(\d{1,2})\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+((?:19|20)\d{2}|\d{2})\b",
    re.IGNORECASE,
)

BIRTH_HINTS = ("г.р", "г.р.", "год рождения", "дата рождения", "родил", "родилась", "birth")
HISTORY_HINTS = (
    "анамнез", "ранее", "предыду", "в прошлом", "прошл", "контроль",
    "контрольный", "повторить", "предыдущ", "сравнение", "динамик",
    "история", "наблюдение", "дата следующего визита", "следующего визита",
)

CONSULT_ANCHORS = (
    "дата первичного осмотра",
    "дата заключительного осмотра",
    "дата осмотра",
    "дата приема",
    "дата приёма",
    "дата визита",
    "прием врача-маммолога",
    "приём врача-маммолога",
    "консультац",
)

ULTRASOUND_ANCHORS = (
    "узи", "ультразвук", "ультразвуков", "исследовани", "протокол", "заключение",
)

SECOND_DOC_HEADER_RE = re.compile(
    r"(?:^|\n)\s*"
    r"(?:медицинское\s+заключение|"
    r"осмотр\s+врача\s+маммолога(?:,\s*первичный)?|"
    r"при[её]м\s+врача-?маммолога|"
    r"протокол\s+исследования\s+молочных\s+желез|"
    r"сүт\s+бездерін\s+зерттеу\s+хаттамасы|"
    r"протокол(?:\s+узи)?|"
    r"узи\s+молочных\s+желез|"
    r"mindray)\b",
    re.IGNORECASE,
)

SECTION_HEADER_HINTS = (
    "жалобы", "анамнез заболевания", "анамнез жизни", "перенесенные заболевания",
    "перенесенные операции", "перенесенные травмы", "вредные привычки",
    "наличие наследственных заболеваний", "непереносимость", "наследственность",
    "anamnesis specialis", "status localis", "результаты обследования",
    "результаты инструментальных", "заключительный диагноз", "рекомендовано",
    "фио пациента", "дата рождения", "контактный телефон", "цель обследования",
    "оң жақ омырау безі", "сол жақ омырау безі", "регионарные", "ұйғарым",
)

RU_MONTH_FOLDER = {
    "январь": 1, "февраль": 2, "март": 3, "апрель": 4, "май": 5, "июнь": 6,
    "июль": 7, "август": 8, "сентябрь": 9, "октябрь": 10, "ноябрь": 11, "декабрь": 12,
}


# ─────────────────────────────────────────────
# Text normalization
# ─────────────────────────────────────────────
def normalize_date_like_text(s: str) -> str:
    if not s:
        return s
    s = s.replace("\xa0", " ")
    s = re.sub(r"(?<=\d),(?=\d)", ".", s)
    s = re.sub(r"(?<=\d)\s+([./\-])\s*(?=\d)", r"\1", s)
    s = re.sub(r"(?<=\d)[./\-]\s+(?=\d)", ".", s)
    s = re.sub(r"(?<=\d)\s+(?=\d{4}\b)", "", s)
    s = re.sub(r"(?<=\d)\.\-(?=\d)", ".", s)
    s = re.sub(r"(?<=\d)\./(?=\d)", ".", s)
    s = re.sub(r"(?<=\d)/\.(?=\d)", "/", s)
    s = re.sub(r"(?<=\d)\.\.(?=\d)", ".", s)
    s = re.sub(r"(?<=\d)//(?=\d)", "/", s)
    s = re.sub(r"(?<=\d)/(?=\d{1,2}\.\d{4}\b)", ".", s)
    s = re.sub(r"(?<=\d)\.(?=\d{1,2}/\d{4}\b)", "/", s)
    s = re.sub(r"[ ]{2,}", " ", s)
    return s


def normalize_text_for_parsing(text: str) -> str:
    if not text:
        return text
    t = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    t = normalize_date_like_text(t)
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def line_looks_like_section_header(line: str) -> bool:
    s = (line or "").strip().lower().strip(":")
    if not s:
        return False
    if any(s.startswith(h) for h in SECTION_HEADER_HINTS):
        return True
    return bool(re.match(r"^[а-яa-zё][^\n]{0,50}:$", s, re.IGNORECASE))


def count_doc_meta_markers(text: str) -> int:
    markers = [
        r"(?:^|\n)\s*Дата\s*[:\-]",
        r"(?:^|\n)\s*Время\s*[:\-]",
        r"(?:^|\n)\s*№\s*истории\s*болезни\s*[:\-]",
        r"(?:^|\n)\s*(?:ФИО\s*пациента|Т\.\s*А\.\s*Ө\.\s*\(Ф\.\s*И\.\s*О\.\))\s*[:\-]",
        r"(?:^|\n)\s*(?:Дата\s*рождения|Туған\s+күні\s*\(Дата\s+рождения\))\s*[:\-]",
    ]
    return sum(1 for pat in markers if re.search(pat, text, re.IGNORECASE))


# ─────────────────────────────────────────────
# Patient identity extraction
# ─────────────────────────────────────────────
def normalize_name(raw_fio: str) -> str:
    s = (raw_fio or "").strip().lower().replace("ё", "е")
    s = s.replace(",", " ").replace(".", " ").replace(";", " ").replace(":", " ")
    s = MULTISPACE_RE.sub(" ", s).strip()
    if not s:
        return ""
    parts = [p for p in s.split(" ") if p]
    cleaned = []
    for p in parts:
        p2 = NON_LETTER_RE.sub("", p).strip("-")
        if p2:
            cleaned.append(p2)
    return " ".join(cleaned).strip()


def clean_fio_candidate(raw_fio: str) -> str:
    s = (raw_fio or "").strip()
    if not s:
        return ""
    s = re.sub(r"[\[\]]", " ", s)
    s = re.sub(r"([0-3]?\d[.\-/][01]?\d[.\-/]\d{4}).*$", "", s).strip(" ,;:-")
    s = MULTISPACE_RE.sub(" ", s).strip()
    return s


def normalize_dob(raw_dob: str) -> str:
    s = normalize_date_like_text((raw_dob or "").strip().lower())
    s = s.replace("г.р.", "").replace("г.р", "").replace("гр.", "").replace("гр", "").strip()
    m = DATE_GENERIC_RE.search(s)
    if not m:
        return ""
    d = int(m.group(1))
    mo = int(m.group(2))
    y_raw = m.group(3)
    y = int(y_raw) if len(y_raw) == 4 else _convert_2digit_year(int(y_raw))
    try:
        return datetime(y, mo, d).strftime("%Y-%m-%d")
    except ValueError:
        return ""


def dob_to_birth_year(raw_dob: str) -> str:
    iso = normalize_dob(raw_dob)
    if iso and len(iso) >= 4:
        year = iso[:4]
        try:
            y = int(year)
            if 1920 <= y <= 2010:
                return year
        except ValueError:
            pass
    return ""


def fio_from_filename(path: Path) -> str:
    stem = path.stem
    s = stem.replace("ё", "е")
    s = re.sub(r"[,_()]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return ""
    tokens = s.split(" ")
    cleaned = []
    for t in tokens:
        tl = t.lower().strip(".-")
        if not tl:
            continue
        if any(sw in tl for sw in FILENAME_STOPWORDS):
            continue
        cleaned.append(t.strip(".-"))
    if len(cleaned) < 2:
        return ""
    surname = cleaned[0]
    t1 = cleaned[1]
    t1_letters = re.sub(r"[^A-Za-zА-Яа-я]", "", t1)
    if t1_letters and len(t1_letters) <= 3 and t1_letters.isalpha():
        initials = t1_letters.upper()
        return f"{surname} {initials}"
    name = cleaned[1] if len(cleaned) >= 2 else ""
    patron = cleaned[2] if len(cleaned) >= 3 else ""
    if not name:
        return ""
    ini1 = re.sub(r"[^A-Za-zА-Яа-яЁё]", "", name)[:1]
    ini2 = re.sub(r"[^A-Za-zА-Яа-яЁё]", "", patron)[:1] if patron else ""
    if not ini1:
        return ""
    initials = (ini1 + ini2).upper()
    return f"{surname} {initials}"


def extract_fio_and_dob(text: str, path: Path) -> tuple[str, str, str]:
    fio = ""
    dob = ""

    m1 = FIO_RE.search(text)
    if m1:
        fio = clean_fio_candidate(m1.group(1))
        embedded_dob = re.search(r"([0-3]?\d[.\-/][01]?\d[.\-/]\d{4})", m1.group(1))
        if embedded_dob:
            dob = embedded_dob.group(1).strip()

    if not dob:
        m2 = DOB_RE.search(text)
        if m2:
            dob = m2.group(1).strip()

    if fio and not dob:
        m3 = DOB_GR_RE.search(text)
        if m3:
            dob = m3.group(1).strip()

    if fio and dob:
        return fio, dob, "docx"

    fio_fn = fio_from_filename(path)
    if fio_fn and dob:
        return fio_fn, dob, "filename"

    return "", "", "none"


def hmac_hex(secret: str, msg: str, length: int = 24) -> str:
    digest = hmac.new(secret.encode("utf-8"), msg.encode("utf-8"), hashlib.sha256).hexdigest()
    return digest[:length]


def make_patient_id(fio_raw: str, dob_raw: str) -> str:
    secret = os.getenv(PATIENT_SECRET_ENV)
    if not secret:
        raise RuntimeError(
            f"Missing env var {PATIENT_SECRET_ENV}. Example:\n"
            f'  export {PATIENT_SECRET_ENV}="..."\n'
            f"Do NOT hardcode the secret in the code."
        )
    fio_norm = normalize_name(fio_raw)
    dob_iso = normalize_dob(dob_raw)
    if not fio_norm or not dob_iso:
        return ""
    key = f"v1|{fio_norm}|{dob_iso}"
    return hmac_hex(secret, key, length=24)


def make_file_id(path: Path) -> str:
    secret = os.getenv(PATIENT_SECRET_ENV)
    if not secret:
        raise RuntimeError(f"Missing env var {PATIENT_SECRET_ENV}")
    try:
        rel = str(path.relative_to(PROJECT_ROOT))
    except Exception:
        rel = str(path)
    return hmac_hex(secret, f"file|v1|{rel}", length=16)


# ─────────────────────────────────────────────
# Date extraction helpers
# ─────────────────────────────────────────────
def _convert_2digit_year(yy: int) -> int:
    return 1900 + yy if yy >= TWO_DIGIT_YEAR_CUTOFF else 2000 + yy


def safe_datetime(day: int, month: int, year: int) -> datetime | None:
    try:
        if year < MIN_YEAR or year > MAX_YEAR:
            return None
        return datetime(year, month, day)
    except Exception:
        return None


def parse_any_date_str(s: str) -> datetime | None:
    s = normalize_date_like_text((s or "").strip().lower())
    if not s:
        return None

    m = DATE_GENERIC_RE.search(s)
    if m:
        day = int(m.group(1))
        month = int(m.group(2))
        year_raw = m.group(3)
        year = int(year_raw) if len(year_raw) == 4 else _convert_2digit_year(int(year_raw))
        return safe_datetime(day, month, year)

    m = RU_MONTH_RE.search(s)
    if m:
        day = int(m.group(1))
        month = RU_MONTHS.get(m.group(2).lower(), 0)
        year_raw = m.group(3)
        if month:
            year = int(year_raw) if len(year_raw) == 4 else _convert_2digit_year(int(year_raw))
            return safe_datetime(day, month, year)

    return None


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                if t:
                    parts.append(t)
    return "\n".join(parts)


def trim_to_main_patient_block(text: str, path: Path) -> str:
    """
    Cut off a foreign trailing block if the same .docx contains
    a second document or another patient.
    """
    if not text:
        return text

    fio_raw, dob_raw, _ = extract_fio_and_dob(text, path)
    main_fio = normalize_name(fio_raw)
    main_dob = normalize_dob(dob_raw)
    cut_positions: list[int] = []

    def add_cut(pos: int) -> None:
        if pos >= 500:
            cut_positions.append(pos)

    for m in PATIENT_NAME_BLOCK_RE.finditer(text):
        candidate_fio = normalize_name(clean_fio_candidate(m.group(1)))
        if main_fio and candidate_fio and candidate_fio != main_fio and m.start() >= 500:
            add_cut(m.start())

    for m in PATIENT_DOB_BLOCK_RE.finditer(text):
        candidate_dob = normalize_dob(m.group(1))
        if main_dob and candidate_dob and candidate_dob != main_dob and m.start() >= 500:
            add_cut(m.start())

    for m in SECOND_DOC_HEADER_RE.finditer(text):
        if m.start() < 500:
            continue

        tail = text[m.start(): m.start() + 2200]
        foreign_name_found = False
        for nm in PATIENT_NAME_BLOCK_RE.finditer(tail):
            candidate_fio = normalize_name(clean_fio_candidate(nm.group(1)))
            if main_fio and candidate_fio and candidate_fio != main_fio:
                foreign_name_found = True
                break

        foreign_dob_found = False
        for dm in PATIENT_DOB_BLOCK_RE.finditer(tail):
            candidate_dob = normalize_dob(dm.group(1))
            if main_dob and candidate_dob and candidate_dob != main_dob:
                foreign_dob_found = True
                break

        meta_count = count_doc_meta_markers(tail)
        looks_like_template_doc = bool(re.search(
            r"(?:^|\n)\s*осмотр\s+врача\s+маммолога(?:,\s*первичный)?\b",
            tail,
            re.IGNORECASE,
        )) and meta_count >= 2

        if foreign_name_found or foreign_dob_found or meta_count >= 3 or looks_like_template_doc:
            add_cut(m.start())

    if not cut_positions:
        return text

    cut_at = min(cut_positions)
    trimmed = text[:cut_at].rstrip()
    if len(trimmed) < 300:
        return text
    return trimmed


def has_any_hint_near(text_lower: str, pos: int, hints: tuple[str, ...], window: int) -> bool:
    start = max(0, pos - window)
    end = min(len(text_lower), pos + window)
    snippet = text_lower[start:end]
    return any(h in snippet for h in hints)


def collect_date_candidates(text: str) -> list[dict]:
    t = normalize_text_for_parsing(text).lower()
    out: list[dict] = []

    for m in DATE_GENERIC_RE.finditer(t):
        day = int(m.group(1))
        month = int(m.group(2))
        year_raw = m.group(3)
        year = int(year_raw) if len(year_raw) == 4 else _convert_2digit_year(int(year_raw))
        dt = safe_datetime(day, month, year)
        if dt is None:
            continue
        if has_any_hint_near(t, m.start(), BIRTH_HINTS, window=55):
            continue
        out.append({"pos": m.start(), "dt": dt, "raw_kind": "generic"})

    for m in RU_MONTH_RE.finditer(t):
        day = int(m.group(1))
        month_name = m.group(2).lower()
        year_raw = m.group(3)
        month = RU_MONTHS.get(month_name, 0)
        if month == 0:
            continue
        year = int(year_raw) if len(year_raw) == 4 else _convert_2digit_year(int(year_raw))
        dt = safe_datetime(day, month, year)
        if dt is None:
            continue
        if has_any_hint_near(t, m.start(), BIRTH_HINTS, window=65):
            continue
        out.append({"pos": m.start(), "dt": dt, "raw_kind": "ru_month"})

    out.sort(key=lambda x: x["pos"])
    deduped: list[dict] = []
    seen = set()
    for item in out:
        key = (item["pos"], item["dt"].date().isoformat())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def score_candidate(text_lower: str, candidate_pos: int, doc_kind: str) -> int:
    score = 0

    if has_any_hint_near(text_lower, candidate_pos, HISTORY_HINTS, window=70):
        score -= 40

    if doc_kind == "consult":
        if has_any_hint_near(text_lower, candidate_pos, CONSULT_ANCHORS, window=90):
            score += 80
        if has_any_hint_near(text_lower, candidate_pos, ("дата первичного осмотра",), window=120):
            score += 60
        if has_any_hint_near(text_lower, candidate_pos, ("дата осмотра",), window=120):
            score += 40
        if has_any_hint_near(text_lower, candidate_pos, ("дата приема", "дата приёма", "дата визита"), window=120):
            score += 35
    else:
        if has_any_hint_near(text_lower, candidate_pos, ULTRASOUND_ANCHORS, window=120):
            score += 80
        if has_any_hint_near(text_lower, candidate_pos, ("узи",), window=140):
            score += 60
        if has_any_hint_near(text_lower, candidate_pos, ("дата исследования", "дата проведения", "дата выполнения"), window=160):
            score += 35

    if candidate_pos < 400:
        score += 10
    if candidate_pos < 200:
        score += 5

    return score


def doc_kind_from_visit_type(visit_type: str) -> str:
    vt = (visit_type or "unknown").strip().lower()
    if vt in {"ultrasound", "ultrasound_onco"}:
        return "ultrasound"
    return "consult"


def _pick_anchor_date(text: str, specs: list[tuple[str, re.Pattern, int]]) -> tuple[str | None, str]:
    t = normalize_text_for_parsing(text).lower()
    for rule, pattern, window in specs:
        for m in pattern.finditer(t):
            tail = t[m.end(): m.end() + window]
            dt = parse_any_date_str(tail)
            if dt is None:
                continue
            if has_any_hint_near(t, m.start(), BIRTH_HINTS, window=70):
                continue
            if has_any_hint_near(t, m.start(), HISTORY_HINTS, window=80):
                continue
            return dt.date().isoformat(), rule
    return None, "none"


def pick_from_priority_patterns(text: str, doc_kind: str) -> tuple[str | None, str]:
    if doc_kind == "ultrasound":
        specs = [
            ("protocol_header_date", re.compile(r"(?:протокол\s+исследования\s+молочных\s+желез|сүт\s+бездерін\s+зерттеу\s+хаттамасы)", re.IGNORECASE), 120),
            ("uzi_from", re.compile(r"узи[^\n]{0,160}от", re.IGNORECASE), 40),
            ("uzi_protocol_from", re.compile(r"(?:протокол[^\n]{0,60})?(?:узи|ультразвук|ультразвуков)[^\n]{0,200}от", re.IGNORECASE), 40),
            ("study_date", re.compile(r"дата\s*(?:проведения|выполнения)?\s*(?:исследовани|узи|ультразвуков)", re.IGNORECASE), 120),
            ("uzi_date_generic", re.compile(r"узи[^\n]{0,160}дата", re.IGNORECASE), 40),
        ]
    else:
        specs = [
            ("primary_exam_date", re.compile(r"дата\s*(?:первичного\s*)?осмотра", re.IGNORECASE), 80),
            ("visit_date", re.compile(r"дата\s*(?:приема|приёма|визита)", re.IGNORECASE), 80),
            ("priem_from", re.compile(r"(?:прием|приём)[^\n]{0,120}от", re.IGNORECASE), 40),
        ]
    return _pick_anchor_date(text, specs)


def choose_main_visit_date(text: str, visit_type: str) -> tuple[str | None, int, str, str]:
    normalized = normalize_text_for_parsing(text)
    doc_kind = doc_kind_from_visit_type(visit_type)

    exact, rule = pick_from_priority_patterns(normalized, doc_kind)
    if exact:
        candidates = collect_date_candidates(normalized)
        return exact, len(candidates), rule, "exact_anchor"

    candidates = collect_date_candidates(normalized)
    if not candidates:
        return None, 0, "none", "none"

    if len(candidates) == 1:
        return candidates[0]["dt"].date().isoformat(), 1, "fallback_single", "fallback_single"

    t = normalized.lower()
    best = None
    best_score = None
    for c in candidates:
        sc = score_candidate(t, c["pos"], doc_kind)
        if best_score is None or sc > best_score:
            best_score = sc
            best = c
        elif sc == best_score and best and c["pos"] < best["pos"]:
            best = c

    return best["dt"].date().isoformat(), len(candidates), "fallback_scored", "fallback_scored"


def infer_visit_type_for_unknown(original_vt: str, doc_kind: str, rule_used: str) -> str:
    vt = (original_vt or "unknown").strip().lower()
    if vt != "unknown":
        return vt
    if rule_used in {"primary_exam_date", "visit_date", "priem_from"}:
        return "consult"
    if rule_used in {"protocol_header_date", "uzi_from", "uzi_protocol_from", "study_date", "uzi_date_generic"}:
        return "ultrasound"
    if rule_used in {"fallback_single", "fallback_scored"}:
        return "ultrasound" if doc_kind == "ultrasound" else "consult"
    return "unknown"


def approx_from_month_year(row: dict) -> str | None:
    month_name = (row.get("month") or "").strip().lower().replace("ё", "е")
    year_str = (row.get("year") or "").strip()
    if not month_name or not year_str:
        return None
    m = RU_MONTH_FOLDER.get(month_name)
    if not m:
        return None
    try:
        y = int(year_str)
    except ValueError:
        return None
    dt = safe_datetime(1, m, y)
    return dt.date().isoformat() if dt else None


# ─────────────────────────────────────────────
# Clinical field extraction
# ─────────────────────────────────────────────
def _first_line_after(text: str, anchors: tuple[str, ...], max_chars: int = 300) -> str:
    normalized = normalize_text_for_parsing(text)
    t_lower = normalized.lower()
    for anchor in anchors:
        pos = t_lower.find(anchor.lower())
        if pos == -1:
            continue
        start = pos + len(anchor)
        while start < len(normalized) and normalized[start] in " :\t-–—":
            start += 1

        chunk = normalized[start: start + max_chars]
        lines = [ln.strip() for ln in chunk.split("\n")]
        kept: list[str] = []
        for idx, line in enumerate(lines):
            if not line:
                if kept:
                    break
                continue
            if idx > 0 and line_looks_like_section_header(line):
                break
            kept.append(line)
            if len(" ".join(kept)) >= max_chars:
                break

        value = " ".join(kept).strip()
        value = MULTISPACE_RE.sub(" ", value)
        return value[:max_chars]
    return ""


def extract_clinical_fields(text: str) -> dict:
    text = normalize_text_for_parsing(text)
    fields: dict = {}

    m = DOB_RE.search(text)
    if not m:
        m = DOB_GR_RE.search(text)
    if m:
        dob_raw = m.group(1)
        iso = normalize_dob(dob_raw)
        if iso and len(iso) >= 4:
            try:
                y = int(iso[:4])
                fields["birth_year"] = str(y) if 1920 <= y <= 2010 else ""
            except ValueError:
                fields["birth_year"] = ""
        else:
            fields["birth_year"] = ""
    else:
        fields["birth_year"] = ""

    menarche_match = re.search(
        r"менарх[еэ]\s*[сc]?\s*[:\-]?\s*(\d{1,2}(?:[.\-]\d{1,2})?)",
        text, re.IGNORECASE,
    )
    fields["menarche_age_raw"] = menarche_match.group(1).strip() if menarche_match else ""

    preg_m = re.search(r"берем[её]нност[иь]\s*[-–—:]\s*(\d+|нет|-)", text, re.IGNORECASE)
    fields["pregnancies_raw"] = preg_m.group(1).strip() if preg_m else ""

    births_m = re.search(r"родов\s*[-–—:]\s*(\d+|нет|-)", text, re.IGNORECASE)
    fields["births_raw"] = births_m.group(1).strip() if births_m else ""

    abort_m = re.search(r"аборт[ао]в\s*[-–—:]\s*(\d+|нет|-)", text, re.IGNORECASE)
    fields["abortions_raw"] = abort_m.group(1).strip() if abort_m else ""

    misc_m = re.search(r"(?:с[/\\]п\s*)?выкидыш[еэей]{0,2}\s*[-–—:]\s*(\d+|нет|-)", text, re.IGNORECASE)
    fields["miscarriages_raw"] = misc_m.group(1).strip() if misc_m else ""

    bf_m = re.search(r"корм(?:ила|ление)[^:\n]{0,20}(?:грудью)?\s*[:\-–—]?\s*([^\n]{1,60})", text, re.IGNORECASE)
    fields["breastfed_raw"] = bf_m.group(1).strip() if bf_m else ""

    compl_m = re.search(
        r"осложнени[яей][^:\n]{0,30}(?:кормлени[яей]|лактац)[^:\n]{0,20}[:\-–—]?\s*([^\n]{1,150})",
        text,
        re.IGNORECASE,
    )
    if not compl_m:
        compl_m = re.search(r"((?:мастит|лактостаз)[^\n]{0,120})", text, re.IGNORECASE)
    fields["breastfed_complications_raw"] = compl_m.group(1).strip() if compl_m else ""

    complaints = _first_line_after(text, ("Жалобы на", "Жалобы:", "Жалобы", "Жалоб"), max_chars=400)
    fields["complaints_raw"] = complaints[:400]

    diagnosis = _first_line_after(
        text,
        (
            "Заключительный диагноз:",
            "Заключительный диагноз",
            "Основной диагноз:",
            "Основной диагноз",
            "Предварительный диагноз:",
            "диагноз:",
        ),
        max_chars=300,
    )
    fields["diagnosis_raw"] = diagnosis[:300]

    icd_m = re.search(r"\b([A-Z]\d{2}(?:\.\d{1,2})?)\b", text)
    fields["icd10_code"] = icd_m.group(1).strip() if icd_m else ""

    birads_m = re.search(r"(?:bi[-\s]?rads|бирадс|birads)\s*[:\-]?\s*([\d|/\\\s]{1,10})", text, re.IGNORECASE)
    fields["birads_raw"] = birads_m.group(1).strip() if birads_m else ""

    nodular = _first_line_after(text, ("Узловые образования", "Объемные образования", "Узловое образование"), max_chars=200)
    fields["nodular_formation_raw"] = nodular[:200]

    heredity = _first_line_after(text, ("Наследственность", "Наследственный анамнез", "Heredity"), max_chars=300)
    fields["heredity_oncology_raw"] = heredity[:300]

    diab_m = re.search(r"сахарн[ыой]{1,2}\s+диабет[а-я]{0,4}[^\n]{0,80}", text, re.IGNORECASE)
    fields["past_diabetes_raw"] = diab_m.group(0).strip() if diab_m else ""

    thyroid = _first_line_after(text, ("Заболевания щитовидной железы", "щитовидная железа", "щитовидной"), max_chars=150)
    fields["past_thyroid_raw"] = thyroid[:150]

    ca_m = re.search(r"[сc][аa][-\s]?15[.\-]?3\s*[-–—:]\s*([\d.,\-–— ]{1,20})\s*(?:ед|u|ме)?", text, re.IGNORECASE)
    fields["ca153_raw"] = ca_m.group(1).strip() if ca_m else ""

    uzi = _first_line_after(text, ("УЗИ молочных желез от", "УЗИ молочных желез", "Ультразвуковое исследование"), max_chars=300)
    fields["uzi_result_raw"] = uzi[:300]

    return fields


# ─────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────
def main() -> None:
    if not IN_CSV.exists():
        raise FileNotFoundError(f"Input CSV not found: {IN_CSV}")

    rows_out: list[dict] = []

    stats_source_before = Counter()
    stats_source_after = Counter()
    stats_quality = Counter()
    stats_rule = Counter()
    stats_errors = Counter()
    stats_trimmed = Counter()

    written = 0
    skipped_no_patient = 0
    skipped_bad_patient = 0

    processed = 0
    parsed_docx = 0
    found_exact_day = 0
    used_month_only = 0
    used_mtime = 0

    fio_source_stats = Counter()

    with IN_CSV.open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        if not reader.fieldnames:
            raise ValueError("Input CSV has no header")

        for i, row in enumerate(reader, 1):
            processed += 1
            if i % PROGRESS_EVERY == 0:
                print(
                    f"Processed {i} rows... docx:{parsed_docx} written:{written} "
                    f"skipped_no_patient:{skipped_no_patient}"
                )

            path = Path(row.get("path", ""))

            original_vt = (row.get("visit_type") or "unknown").strip().lower()
            before_source = (row.get("date_source") or "path").strip() or "path"
            stats_source_before[before_source] += 1

            if not path.exists():
                stats_errors["missing_file"] += 1
                skipped_no_patient += 1
                continue

            try:
                parsed_docx += 1
                raw_text = docx_text(path)
                trimmed_text = trim_to_main_patient_block(raw_text, path)
                if trimmed_text != raw_text:
                    stats_trimmed["trimmed_foreign_tail"] += 1
                text = normalize_text_for_parsing(trimmed_text)
            except Exception:
                stats_errors["docx_read_error"] += 1
                skipped_no_patient += 1
                continue

            fio_raw, dob_raw, fio_source = extract_fio_and_dob(text, path)
            if not fio_raw or not dob_raw:
                skipped_no_patient += 1
                continue

            fio_source_stats[fio_source] += 1

            try:
                patient_id = make_patient_id(fio_raw, dob_raw)
                file_id = make_file_id(path)
            except Exception:
                skipped_bad_patient += 1
                continue

            if not patient_id:
                skipped_bad_patient += 1
                continue

            visit_date = ""
            date_source = before_source
            date_rule_used = "none"
            date_quality = "none"
            candidates_count = 0

            try:
                best_date, cnt, rule_used, quality = choose_main_visit_date(text, original_vt)
                candidates_count = cnt
                date_rule_used = rule_used
                date_quality = quality

                if best_date:
                    visit_date = best_date
                    date_source = "docx"
                    found_exact_day += 1
                else:
                    approx = approx_from_month_year(row)
                    if approx:
                        visit_date = approx
                        date_source = before_source
                        date_quality = "month_only"
                        used_month_only += 1
                    else:
                        mtime = (row.get("mtime") or "").strip()
                        if mtime:
                            visit_date = mtime.split("T")[0]
                            date_source = "mtime"
                            date_quality = "mtime"
                            used_mtime += 1
                        else:
                            date_quality = "none"
            except Exception:
                stats_errors["date_extract_error"] += 1

            vt_inferred = original_vt
            vt_final = original_vt
            if original_vt == "unknown":
                vt_inferred = infer_visit_type_for_unknown(
                    original_vt, doc_kind_from_visit_type(original_vt), date_rule_used
                )
                vt_final = vt_inferred

            try:
                clinical = extract_clinical_fields(text)
            except Exception:
                stats_errors["clinical_extract_error"] += 1
                clinical = {
                    "birth_year": "", "menarche_age_raw": "",
                    "pregnancies_raw": "", "births_raw": "",
                    "abortions_raw": "", "miscarriages_raw": "",
                    "breastfed_raw": "", "breastfed_complications_raw": "",
                    "complaints_raw": "", "diagnosis_raw": "",
                    "icd10_code": "", "birads_raw": "",
                    "nodular_formation_raw": "", "heredity_oncology_raw": "",
                    "past_diabetes_raw": "", "past_thyroid_raw": "",
                    "ca153_raw": "", "uzi_result_raw": "",
                }

            out_row = {}
            for k, v in row.items():
                if k.lower() in {"path", "filename"}:
                    continue
                out_row[k] = v

            out_row["file_id"] = file_id
            out_row["patient_id"] = patient_id
            out_row["fio_source"] = fio_source
            out_row["visit_date"] = visit_date
            out_row["date_source"] = date_source
            out_row["date_rule_used"] = date_rule_used
            out_row["date_quality"] = date_quality
            out_row["date_candidates_count"] = str(candidates_count)
            out_row["visit_type_inferred"] = vt_inferred
            out_row["visit_type_final"] = vt_final
            out_row.update(clinical)

            rows_out.append(out_row)
            written += 1

            stats_source_after[date_source] += 1
            stats_quality[date_quality] += 1
            stats_rule[date_rule_used] += 1

    if not rows_out:
        print("No rows written. Check PATIENT_ID_SECRET and patterns.")
        return

    fieldnames = list(rows_out[0].keys())
    with OUT_CSV.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        w.writerows(rows_out)

    print("\nDone.")
    print("Processed rows:", processed)
    print("Docx parsed:", parsed_docx)
    print("Rows written:", written)
    print("Skipped (no FIO/DOB):", skipped_no_patient)
    print("Skipped (bad normalize/HMAC):", skipped_bad_patient)
    print("FIO source:", dict(fio_source_stats))
    print("Exact day from docx:", found_exact_day)
    print("Month-only approximations:", used_month_only)
    print("Used mtime fallback:", used_mtime)
    print("Saved:", OUT_CSV)

    print("\nDate source before:")
    for k, v in stats_source_before.most_common():
        print(f"  {k}: {v}")

    print("\nDate source after:")
    for k, v in stats_source_after.most_common():
        print(f"  {k}: {v}")

    print("\nDate quality:")
    for k, v in stats_quality.most_common():
        print(f"  {k}: {v}")

    print("\nRule used (overall):")
    for k, v in stats_rule.most_common():
        print(f"  {k}: {v}")

    if stats_trimmed:
        print("\nTrimmed foreign tails:")
        for k, v in stats_trimmed.most_common():
            print(f"  {k}: {v}")

    if stats_errors:
        print("\nErrors (safe counts):")
        for k, v in stats_errors.most_common():
            print(f"  {k}: {v}")


if __name__ == "__main__":
    main()